import os
from flask import Flask, request, render_template, jsonify, send_file
import google.generativeai as genai
import time
import shutil
import io

# For Word document generation
from docx import Document


app = Flask(__name__)

# --- Gemini API Key Configuration ---
# WARNING: Storing API keys directly in code is NOT recommended for production environments.
# For local development and personal use, it might be convenient.
# For deployment, ALWAYS use environment variables or a secure secret management system.
YOUR_GEMINI_API_KEY = os.environ.get("gemini_key") # <--- THIS IS THE CRUCIAL CHANGE

genai.configure(api_key=YOUR_GEMINI_API_KEY)

if YOUR_GEMINI_API_KEY == "YOUR_ACTUAL_GEMINI_API_KEY_HERE" or YOUR_GEMINI_API_KEY.strip() == "":
    print("\n--- ERROR: Gemini API Key Not Set in Code! ---")
    print("Please replace 'YOUR_ACTUAL_GEMINI_API_KEY_HERE' with your actual Gemini API key in app.py.")
    exit("Gemini API Key Missing. Exiting application.")

@app.route('/')
def index():
    """Renders the main web page for audio upload and transcription."""
    return render_template('index.html')

@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    """Handles the audio file upload, performs dual-language transcription, and returns the result."""
    if 'audio' not in request.files:
        return jsonify({'error': 'No audio file provided. Please select a file.'}), 400

    audio_file_storage = request.files['audio']
    if audio_file_storage.filename == '':
        return jsonify({'error': 'No audio file selected. Please choose a file.'}), 400

    temp_dir = None
    audio_part = None
    temp_input_audio_path = None

    try:
        request_id = int(time.time()) 
        temp_dir_base = "temp_transcriptions"
        os.makedirs(temp_dir_base, exist_ok=True)
        temp_dir = os.path.join(temp_dir_base, f"request_{request_id}")
        os.makedirs(temp_dir) 
        print(f"Created temporary directory for request: {temp_dir}")

        original_filename = audio_file_storage.filename
        extension = os.path.splitext(original_filename)[1].lower()
        if not extension:
            extension = ".wav" 
        
        temp_input_audio_path = os.path.join(temp_dir, f"input_audio{extension}")
        audio_file_storage.save(temp_input_audio_path)
        print(f"Uploaded audio saved locally to: {temp_input_audio_path}")

        print("Attempting to transcribe the entire audio in a single API request using Gemini 2.5 Pro.")

        audio_part = genai.upload_file(temp_input_audio_path)
        print(f"  Audio uploaded to Gemini Files API. File name: {audio_part.name}, current state: {audio_part.state.name}")

        print("  Waiting for Gemini Files API to process audio (can take several minutes for long files)...")
        while audio_part.state.name == "PROCESSING":
            print('.', end='', flush=True)
            time.sleep(5)
            audio_part = genai.get_file(name=audio_part.name)
        print("\n")

        if audio_part.state.name != "ACTIVE":
            error_message = (f"Audio file processing failed or is not ready ('ACTIVE') on Gemini's servers. "
                             f"Final state: {audio_part.state.name}. "
                             f"This could be due to an unsupported audio format, file corruption, "
                             f"or an internal Gemini Files API issue. Please ensure your audio is clean, "
                             f"in a common format (e.g., MP3, WAV), and try again.")
            raise Exception(error_message)

        model = genai.GenerativeModel('gemini-2.5-pro') 

        # --- Bengali Transcription Prompt ---
        prompt_bengali = """
        This is an audio recording of a qualitative research interview (either a Key Informant Interview (KII) or In-Depth Interview (IDI)).
        Please transcribe the entire audio into Bengali, including very detailed timestamps and speaker identification.
        For speaker identification, use "Speaker A", "Speaker B", "Speaker C", etc.
        Include timestamps at the beginning of each speaker's turn and whenever there's a significant pause or change in topic.
        Also, provide minute-level timestamps at the start of each new minute mark (e.g., [00:01:00], [00:02:00]).
        Pay extremely close attention to non-verbal cues and emotional states. If a speaker is crying, indicate it with "[কাঁদছে]" (crying). If they are smiling or laughing, indicate it with "[হাসছে]" (smiling/laughing). Also note other significant sounds like "[কাশি]" (coughing), "[নীরবতা]" (silence), "[অন্যান্য শব্দ]" (other sounds).
        Ensure the transcription captures every single detail and nuance relevant for qualitative analysis.

        Example output format:
        [00:00:05] Speaker A: আপনি কেমন আছেন?
        [00:00:10] Speaker B: আমি ভালো আছি, ধন্যবাদ। [হাসছে]
        [00:00:15] Speaker A: আপনার গবেষণার বিষয় কি?
        [00:00:22] Speaker B: (কিছুক্ষণ নীরবতা) এটি একটি জটিল বিষয়। [কাঁদছে]
        [00:01:00] Speaker A: পরবর্তী প্রশ্ন...
        """

        # --- English Transcription/Translation Prompt (ENHANCED for Contextualization & Storytelling) ---
        prompt_english = """
        This is an audio recording of a qualitative research interview (either a Key Informant Interview (KII) or In-Depth Interview (IDI)) conducted in Bengali.
        Your task is to transcribe and translate the entire audio into **highly contextualized, natural-sounding English**, reflecting the conversational and storytelling style typical of in-depth qualitative interviews.

        **Key Requirements:**
        1.  **Narrative Flow & Naturalness:** Ensure the English translation reads coherently and naturally as spoken language, avoiding overly literal word-for-word translation. Prioritize idiomatic expressions and natural phrasing that captures the speaker's original intent and emotional tone.
        2.  **Contextual Understanding:** Analyze the entire audio to understand the broader context of the conversation. Use this understanding to contextualize responses, making the English transcription a fluid and analytically useful narrative for qualitative research.
        3.  **Detailed Timestamps:** Include very precise timestamps at the beginning of each speaker's turn, significant pauses, or topic shifts. Also, include minute-level timestamps (e.g., [00:01:00], [00:02:00]).
        4.  **Speaker Identification:** Use clear labels like "Speaker A", "Speaker B", "Speaker C", etc.
        5.  **Non-Verbal & Emotional Cues:** Pay close attention to and transcribe relevant non-verbal cues and emotional states in English, such as:
            * [Crying] or [Sobbing]
            * [Smiling] or [Laughing]
            * [Coughing]
            * [Silence] or [Pause]
            * [Other sounds] (e.g., [Door opens], [Phone rings])

        The aim is to provide a rich, narrative English transcript that facilitates deep qualitative analysis, preserving the essence and nuance of the original Bengali interview.

        Example output format:
        [00:00:05] Speaker A: How are you doing today?
        [00:00:10] Speaker B: I'm doing quite well, thank you very much. [Smiling]
        [00:00:15] Speaker A: Could you elaborate on your research topic?
        [00:00:22] Speaker B: (A brief silence) Well, it's quite a complex subject, actually. [Crying]
        [00:01:00] Speaker A: Moving on to the next question...
        """
        
        print("Sending Bengali transcription request to Gemini API...")
        response_bengali = model.generate_content([prompt_bengali, audio_part])
        print("Bengali transcription received.")

        print("Sending English transcription request to Gemini API...")
        response_english = model.generate_content([prompt_english, audio_part])
        print("English transcription received.")
        
        # Return both transcriptions in the JSON response
        return jsonify({
            'bengali_transcription': response_bengali.text,
            'english_transcription': response_english.text
        })

    except Exception as e:
        error_message = f"An error occurred during transcription: {str(e)}"
        print(f"ERROR: {error_message}")
        return jsonify({'error': error_message}), 500
    finally:
        # --- Cleanup Section ---
        if temp_input_audio_path and os.path.exists(temp_input_audio_path):
            try:
                os.remove(temp_input_audio_path)
                print(f"Cleaned up local temporary input audio file: {temp_input_audio_path}")
            except Exception as cleanup_e:
                print(f"WARNING: Error deleting temporary input audio file {temp_input_audio_path}: {cleanup_e}")

        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                print(f"Cleaned up local temporary directory: {temp_dir}")
            except Exception as cleanup_e:
                print(f"WARNING: Error cleaning up temporary directory {temp_dir}: {cleanup_e}")
        
        if audio_part and audio_part.name:
            try:
                time.sleep(2) # Small delay for eventual consistency before deletion
                genai.delete_file(audio_part.name)
                print(f"Deleted Gemini API temporary file: {audio_part.name}")
            except Exception as delete_e:
                print(f"WARNING: Error deleting Gemini API file {audio_part.name}: {delete_e}")
    
    return jsonify({'error': 'An unexpected error occurred during file handling or an unhandled state.'}), 500

# --- New Route for Bengali Word Document Download ---
@app.route('/download_bengali_docx', methods=['POST'])
def download_bengali_docx():
    """Receives Bengali transcription text and generates a Word document for download."""
    bengali_text = request.json.get('transcription', '')
    
    if not bengali_text:
        return jsonify({'error': 'No Bengali transcription text provided for download.'}), 400

    document = Document()
    document.add_heading('Bengali Transcription (বাংলা প্রতিলিপি) - Qualitative Interview', level=1)
    
    lines_bengali = bengali_text.split('\n')
    for line in lines_bengali:
        if line.strip(): # Avoid adding empty paragraphs for blank lines
            document.add_paragraph(line.strip())

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0) # Rewind the stream to the beginning for reading

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        download_name='bengali_transcription.docx',
        as_attachment=True # Forces download rather than displaying in browser
    )

# --- New Route for English Word Document Download ---
@app.route('/download_english_docx', methods=['POST'])
def download_english_docx():
    """Receives English transcription text and generates a Word document for download."""
    english_text = request.json.get('transcription', '')
    
    if not english_text:
        return jsonify({'error': 'No English transcription text provided for download.'}), 400

    document = Document()
    document.add_heading('English Transcription - Qualitative Interview', level=1)
    
    lines_english = english_text.split('\n')
    for line in lines_english:
        if line.strip():
            document.add_paragraph(line.strip())

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        download_name='english_transcription.docx',
        as_attachment=True
    )
# --- End New Routes ---

if __name__ == '__main__':
    app.run(debug=True, port=5000)